from docx import Document

def iter_from_file(doc, joyo=False):
    document = Document(doc)
    if joyo:
        paragraphs = document.paragraphs[:-5]
    else:
        paragraphs = document.paragraphs
    
    list = ""
    for paragraph in paragraphs:
        for elt in paragraph.iter_inner_content():
            if elt.text not in "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789 -_éèàçù":
                list += elt.text
    
    return iter(list)

def pattern_from_img(img):
    pattern = []
    *list_pix, = img.getdata()
    width, height = img.size
    pixels = [list_pix[i * width:(i+1)*width] for i in range(height)]
    for row in pixels:
        line = []
        val = 0
        counter = 0
        for elt in row:
            if elt == val:
                counter += 1
            else:
                line += counter,
                val = elt
                counter = 1
        
        line += counter,
        pattern += line,
    return pattern
    
    
def register(name, img, pattern):
    pass